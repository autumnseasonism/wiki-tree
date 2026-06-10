#!/usr/bin/env python3
"""
convert_documents.py — 将多种格式文档统一转换为 Markdown
用法: python convert_documents.py --scan-report /path/to/scan.json --output /path/to/output

可选依赖（按需，核心 md/txt/json/csv 转换纯标准库）: pip install python-docx  # 仅 .docx
                                                  pip install PyMuPDF      # 仅 .pdf
"""

import os
import sys
import json
import hashlib
import argparse
from pathlib import Path
from datetime import datetime, timezone


def convert_word(file_path: str) -> str:
    """Word (.docx) → Markdown"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("缺少依赖 python-docx，无法转换 .docx：pip install python-docx")

    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style_name = para.style.name.lower() if para.style else ""

        if "heading 1" in style_name or "标题 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name or "标题 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name or "标题 3" in style_name:
            lines.append(f"### {text}")
        elif "heading 4" in style_name or "标题 4" in style_name:
            lines.append(f"#### {text}")
        elif "list" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # 表格：追加在段落之后（防御性——失败不影响段落提取）；序号按实际渲染的表连续编号（跳过空表不留空号）
    try:
        tn = 0
        for table in doc.tables:
            trows = [[(c.text or "").strip().replace("|", "\\|") for c in row.cells]
                     for row in table.rows]
            trows = [r for r in trows if any(r)]
            if not trows:
                continue
            tn += 1
            w = max(len(r) for r in trows)
            pad = lambda r: r + [""] * (w - len(r))
            lines.append("")
            lines.append(f"<!-- 表格 {tn} -->")
            lines.append("| " + " | ".join(pad(trows[0])) + " |")
            lines.append("| " + " | ".join(["---"] * w) + " |")
            for r in trows[1:]:
                lines.append("| " + " | ".join(pad(r)) + " |")
    except Exception as _table_err:
        print(f"警告: docx 表格提取失败（段落内容仍保留）: {_table_err}", file=sys.stderr)

    return "\n".join(lines)


def convert_pdf(file_path: str) -> str:
    """PDF → Markdown"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("缺少依赖 PyMuPDF，无法转换 .pdf：pip install PyMuPDF")

    doc = fitz.open(file_path)
    # 加密 PDF：get_text 会返回空串，若不在此报错会被当成"空内容"而非明确失败
    if doc.needs_pass:
        doc.close()
        raise RuntimeError("PDF 已加密（需要密码），无法提取文本")
    pages = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        if text.strip():
            pages.append(f"<!-- 第 {page_num + 1} 页 -->\n{text.strip()}")

    doc.close()
    return "\n\n".join(pages)


def convert_json(file_path: str) -> str:
    """JSON → Markdown"""
    try:
        # utf-8-sig：带 BOM 的 JSON（Windows 工具导出常见）json.load 会直接失败；无 BOM 时等价 utf-8
        with open(file_path, "r", encoding="utf-8-sig") as f:
            data = json.load(f)
    except (json.JSONDecodeError, UnicodeDecodeError) as e:
        raise RuntimeError(f"JSON 解析失败: {e}") from e

    result = json_to_markdown(data, level=2)
    # 与 convert_text 同构的总量截断：JSON_LIST_CAP 只封顶数组项数，
    # 巨型标量/大字典（单键塞日志、base64、聊天导出常见）仍可能产出巨型 md
    if len(result) > MAX_TEXT_CHARS:
        result = (result[:MAX_TEXT_CHARS]
                  + f"\n\n<!-- 文件过大，仅转换前 {MAX_TEXT_CHARS} 字符 -->")
    return result


JSON_LIST_CAP = 1000  # JSON 数组项数上限；超出截断并注明（导出类 JSON 动辄数万条，撑爆下游抽取）


def json_to_markdown(obj, level=2, prefix="") -> str:
    """递归将 JSON 转为 Markdown"""
    lines = []
    heading = "#" * level

    if isinstance(obj, dict):
        for key, value in obj.items():
            if isinstance(value, (dict, list)):
                lines.append(f"{heading} {prefix}{key}")
                lines.append(json_to_markdown(value, level + 1, f"{prefix}{key}."))
            else:
                lines.append(f"**{key}**: {value}")
        lines.append("")
    elif isinstance(obj, list):
        for i, item in enumerate(obj[:JSON_LIST_CAP]):
            if isinstance(item, (dict, list)):
                lines.append(f"{heading} [{i}]")
                lines.append(json_to_markdown(item, level + 1, f"{prefix}[{i}]."))
            else:
                lines.append(f"- {item}")
        if len(obj) > JSON_LIST_CAP:
            lines.append(f"<!-- 列表过大：共 {len(obj)} 项，仅转换前 {JSON_LIST_CAP} 项 -->")
        lines.append("")
    else:
        lines.append(str(obj))

    return "\n".join(lines)


MAX_TEXT_CHARS = 2_000_000  # 纯文本读取上限（字符）；百 MB 级 log 整读会撑爆内存与下游抽取，超出截断并注明


def convert_text(file_path: str) -> str:
    """纯文本 → Markdown"""
    truncated = False
    try:
        # utf-8-sig：剥掉可能的 BOM；无 BOM 时等价 utf-8
        with open(file_path, "r", encoding="utf-8-sig") as f:
            content = f.read(MAX_TEXT_CHARS)
            truncated = bool(f.read(1))
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="gbk") as f:
                content = f.read(MAX_TEXT_CHARS)
                truncated = bool(f.read(1))
        except UnicodeDecodeError:
            raise RuntimeError("无法解码文件（尝试了 UTF-8 和 GBK）")

    # 简单的段落检测：连续非空行合并为段落
    paragraphs = []
    current = []
    for line in content.split("\n"):
        if line.strip():
            current.append(line.strip())
        else:
            if current:
                paragraphs.append(" ".join(current))
                current = []
    if current:
        paragraphs.append(" ".join(current))

    result = "\n\n".join(paragraphs)
    if truncated:
        result += f"\n\n<!-- 文件过大，仅转换前 {MAX_TEXT_CHARS} 字符 -->"
    return result


CSV_ROW_CAP = 1000   # CSV 数据行上限；超出则截断并注明（避免超大表撑爆下游抽取）
CSV_CELL_CAP = 2000  # 单元格字符上限；超出则截断并注明（单格塞日志/base64 时与行截断同一哲学）


def _md_cell(value) -> str:
    """转义单元格内容，使其安全嵌入 Markdown 表格；超长单元格截断并注明。"""
    s = str(value)
    if len(s) > CSV_CELL_CAP:
        s = f"{s[:CSV_CELL_CAP]} …（单元格过长，已截断，原 {len(s)} 字符）"
    return (s.replace("\\", "\\\\").replace("|", "\\|")
            .replace("\r", " ").replace("\n", " ").strip())


def convert_csv(file_path: str) -> str:
    """CSV → Markdown 表格（首行作表头；超大表截断并注明）。"""
    import csv
    # 默认 field_size_limit=131072，单元格超长会抛 csv.Error 整文件失败；
    # 放宽到平台上限（Windows 的 C long 为 32 位，直接给 sys.maxsize 会 OverflowError）
    csv.field_size_limit(min(sys.maxsize, 2**31 - 1))
    rows = None
    last_err = None
    for enc in ("utf-8-sig", "utf-8", "gbk"):
        try:
            with open(file_path, "r", encoding=enc, newline="") as f:
                rows = list(csv.reader(f))
            break
        except (UnicodeDecodeError, csv.Error) as e:
            last_err = e
            continue
    if rows is None:
        raise RuntimeError(f"无法解析 CSV（尝试了 UTF-8 和 GBK）: {last_err}")
    rows = [r for r in rows if any((c or "").strip() for c in r)]  # 丢弃全空行
    if not rows:
        return ""
    total_data = len(rows) - 1
    note = ""
    if total_data > CSV_ROW_CAP:
        rows = [rows[0]] + rows[1:CSV_ROW_CAP + 1]
        note = f"\n\n<!-- 表格过大：共 {total_data} 行数据，仅转换前 {CSV_ROW_CAP} 行 -->"
    width = max(len(r) for r in rows)
    pad = lambda r: list(r) + [""] * (width - len(r))
    lines = ["| " + " | ".join(_md_cell(c) for c in pad(rows[0])) + " |",
             "| " + " | ".join(["---"] * width) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(_md_cell(c) for c in pad(r)) + " |")
    return "\n".join(lines) + note


CONVERTERS = {
    "word": convert_word,
    "pdf": convert_pdf,
    "markdown": None,  # 直接读取
    "json": convert_json,
    "text": convert_text,
    "csv": convert_csv,
}


def _existing_is_same_source(md_path: Path, source_path: str) -> bool:
    """同名 .md 是否就是本源文件（front-matter 里的 source_path 一致）。

    增量重跑/上次中断遗留时用于判断：是 → 覆写既有文件，而不是再生成 -1 重复副本。
    不同源文件恰好同名时（如两个目录各有 report.docx）仍走 -1 区分，互不影响。
    """
    try:
        # 2000：JSON 引号格式下反斜杠翻倍，长 Windows 路径可能超出旧的 800 截断窗口
        head = md_path.read_text(encoding="utf-8")[:2000]
    except (OSError, UnicodeDecodeError):
        return False
    # 整行匹配（带行尾 \n），避免"查询路径恰为已存路径的行前缀"的假阳性
    # （如 /a/report 误命中存有 /a/report.docx 的卡片）。
    # 新产物 source_path 是 JSON 字符串（见 process_file），旧 vault 的裸值行也要认
    quoted = f"source_path: {json.dumps(source_path, ensure_ascii=False)}\n"
    return quoted in head or f"source_path: {source_path}\n" in head


def _demote_md_frontmatter(text: str) -> str:
    """源 .md 若以 YAML front-matter 开头，将其降级为正文顶部的引用块。

    转换器会在每篇输出顶部注入自己的 front-matter；若源文件本身也有 front-matter，
    旧行为是直接拼在其下 → 形成"双 front-matter"，第二段会被当成正文泄漏（且
    verify_entities 只剥第一段）。这里把源 front-matter 降级为带标注的引用块：
    既不再有第二个 YAML 头，内容也仍以正文形式保留、可被实体校验匹配到。
    纯行级处理，不依赖 pyyaml；未检测到 front-matter 时原样返回。
    """
    if not text.startswith("---"):
        return text
    lines = text.splitlines(keepends=True)
    if lines[0].strip() != "---":
        return text
    for i in range(1, len(lines)):
        if lines[i].strip() == "---":
            fm = [ln.rstrip("\n") for ln in lines[1:i] if ln.strip()]
            body = "".join(lines[i + 1:]).lstrip("\n")
            if not fm:
                return body
            quoted = "\n".join("> " + ln for ln in fm)
            return ("<!-- 源文档 front-matter（已降级为引用，避免双 front-matter）-->\n"
                    + quoted + "\n\n" + body)
    return text  # 未闭合的 ---，不当 front-matter 处理


def process_file(file_info: dict, output_dir: Path, content_hashes: dict = None) -> dict:
    """处理单个文件，返回结果信息。

    content_hashes 非 None 时启用文档级内容去重：内容完全相同的文件只保留第一份。
    """
    source_path = file_info["path"]
    category = file_info["category"]
    name = Path(source_path).stem

    # 输出文件名：去重 + 清理
    safe_name = name.replace(" ", "-").replace("/", "-").replace("\\", "-")
    output_path = output_dir / f"{safe_name}.md"

    # 避免重名；但若同名文件正是本源文件（增量重跑/上次中断遗留）→ 覆写，不再生成 -1 副本
    counter = 1
    while output_path.exists():
        if _existing_is_same_source(output_path, source_path):
            break
        output_path = output_dir / f"{safe_name}-{counter}.md"
        counter += 1

    try:
        if category == "markdown":
            try:
                # utf-8-sig：带 BOM 的源 .md 若按 utf-8 读，BOM 会挡在 "---" 前，
                # 下方 front-matter 降级检测失效 → 双 front-matter 泄漏进正文
                content = Path(source_path).read_text(encoding="utf-8-sig")
            except UnicodeDecodeError:
                content = Path(source_path).read_text(encoding="gbk")
            # 源 .md 若自带 front-matter，降级为正文引用块，避免与下方注入的
            # front-matter 形成"双 front-matter"（第二段会泄漏进正文）
            content = _demote_md_frontmatter(content)
        else:
            converter = CONVERTERS.get(category)
            if not converter:
                return {
                    "status": "skipped",
                    "reason": f"不支持的类型: {category}",
                    "source": source_path,
                }
            content = converter(source_path)

        # 空内容不写盘：写出只剩 front-matter 的空 .md 会被登记 done，内容永久丢失且无人发现
        if not content.strip():
            reason = "未提取到文本，疑似扫描版 PDF/空文件"
            # 已入库源后来被清空：同源旧产物不该继续被 reduce/索引，删除并标注
            if output_path.exists() and _existing_is_same_source(output_path, source_path):
                try:
                    output_path.unlink()
                    reason += "（旧产物已删除）"
                except OSError as e:
                    reason += f"（旧产物删除失败: {e}）"
            return {
                "status": "empty",
                "reason": reason,
                "source": source_path,
            }

        # 文档级内容去重：内容完全相同的文件只转换第一份
        if content_hashes is not None:
            digest = hashlib.sha256(content.encode("utf-8")).hexdigest()
            if digest in content_hashes:
                return {
                    "status": "skipped",
                    "reason": f"内容重复，等同于 {content_hashes[digest]}",
                    "source": source_path,
                    "duplicate_of": content_hashes[digest],
                    # scan 时刻捕获的源文件 mtime，沿管线传给 manifest 登记（见 update_manifest）
                    "modified_at": file_info.get("modified_at"),
                }
            content_hashes[digest] = str(output_path)

        # 添加 front-matter；source_path 用 JSON 字符串（合法 YAML 标量），
        # 路径含 " #" / ": " 等 YAML 特殊序列时不再截断/破坏整段
        now = datetime.now(timezone.utc).isoformat()
        fm = f"""---
source_type: local_file
source_path: {json.dumps(source_path, ensure_ascii=False)}
source_format: {file_info.get('extension', 'unknown').lstrip('.')}
converted_at: {now}
file_size_bytes: {file_info.get('size_bytes', 0)}
---

"""
        output_path.write_text(fm + content, encoding="utf-8")

        return {
            "status": "success",
            "source": source_path,
            "output": str(output_path),
            "size": len(content),
            # scan 时刻捕获的源文件 mtime，沿管线传给 manifest 登记（消除抽取期间源被改的时间窗）
            "modified_at": file_info.get("modified_at"),
        }

    except Exception as e:
        return {
            "status": "error",
            "reason": str(e),
            "source": source_path,
        }


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass
    parser = argparse.ArgumentParser(description="批量转换文档为 Markdown")
    parser.add_argument("--scan-report", required=True, help="scan_folder.py 生成的 JSON 报告")
    parser.add_argument("--output", "-o", required=True, help="输出目录")
    parser.add_argument("--limit", type=int, default=0, help="最多处理文件数（0=全部）")
    parser.add_argument("--no-dedup-content", action="store_true",
                        help="关闭文档级内容去重（默认开启：内容完全相同的文件只转换第一份）")
    args = parser.parse_args()

    # 读取扫描报告
    with open(args.scan_report, "r", encoding="utf-8") as f:
        report = json.load(f)

    files = report.get("files", [])
    if args.limit > 0:
        files = files[:args.limit]

    output_dir = Path(args.output) / "documents"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 文档级内容去重（默认开启；--no-dedup-content 关闭）
    content_hashes = None if args.no_dedup_content else {}

    results = {
        "total": len(files),
        "success": 0,
        "error": 0,
        "empty": 0,
        "skipped": 0,
        "details": [],
    }

    for i, file_info in enumerate(files):
        result = process_file(file_info, output_dir, content_hashes)
        results["details"].append(result)

        if result["status"] == "success":
            results["success"] += 1
        elif result["status"] == "error":
            results["error"] += 1
        elif result["status"] == "empty":
            results["empty"] += 1
        else:
            results["skipped"] += 1

        # 进度输出
        if (i + 1) % 10 == 0 or i + 1 == len(files):
            print(f"进度: {i + 1}/{len(files)} "
                  f"(成功: {results['success']}, "
                  f"错误: {results['error']}, "
                  f"空内容: {results['empty']}, "
                  f"跳过: {results['skipped']})")

    # 写入转换报告
    report_path = Path(args.output) / "_conversion_report.json"
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\n转换完成！报告: {report_path}")
    print(f"  成功: {results['success']}")
    print(f"  错误: {results['error']}")
    print(f"  空内容(未写盘，疑似扫描版 PDF/空文件): {results['empty']}")
    print(f"  跳过: {results['skipped']}")


if __name__ == "__main__":
    main()
