#!/usr/bin/env python3
"""
wiki-tree 回归测试套件。
运行: python tests/test_skill.py   （建议先 set PYTHONUTF8=1 / export PYTHONUTF8=1）
     或 python -m pytest tests/test_skill.py（经文件末尾的桥接用例收集）
覆盖 10 个脚本 + 3 个接入包模板的核心行为 + 历次修复：
  管线: scan_folder / generate_wiki_structure / convert_documents / update_manifest /
        verify_entities / compute_centrality / assemble_vault / suggest_dedup
  接入: emit_access_bundle / kb_register / templates(kb_query · kb_ingest · kb_mcp_server · kb_hub_server)
  未覆盖（待补）: emit_doc_summaries / check_deps
  文档: [22] 节守卫 SKILL.md / references 的关键契约（体积、点名命令、输出契约字段）
全程用临时目录（接入层用假 HOME 重定向 ~），不触碰真实 registry / CLAUDE.md / 用户数据；任一断言失败则退出码 1。
每个编号节经 sect() 隔离：一节崩溃记 CRASH 计 FAIL，其余节继续执行，汇总始终打印。
run() 默认把子进程非零退出码计为 FAIL（预期失败的调用传 ok=False）。
可选依赖缺失的用例记 SKIP 并计入汇总；CI 用环境变量 EXPECT_SKIPS=N 断言预期跳过数
（裸环境=3: docx/pdf/hub 进程内；全量=1: kb_mcp_server 缺 SDK 路径），
防止「裸环境腿误装了依赖 / 全量腿漏装」这类环境前置条件静默失效。
"""
import io
import os, re, sys, json, time, subprocess, tempfile, shutil
import contextlib
from pathlib import Path

SC = Path(__file__).resolve().parent.parent / "scripts"
SKILL_ROOT = SC.parent
PY = sys.executable
fails = []
skips = []

def check(name, cond, extra=""):
    print(("  PASS " if cond else "  FAIL ") + name + (f"  [{extra}]" if extra and not cond else ""))
    if not cond:
        fails.append(name)

def skip(name):
    print("  SKIP " + name)
    skips.append(name)

@contextlib.contextmanager
def sect(title):
    """分节隔离：一节崩溃只损失本节（计 1 个 FAIL），后续节照常执行。"""
    print(title)
    try:
        yield
    except Exception as e:  # noqa: BLE001
        msg = f"{title} crashed: {type(e).__name__}: {e}"
        print("  CRASH " + msg)
        fails.append(msg)

def run(*a, env=None, ok=True):
    """跑脚本子进程。ok=True 时非零退出码本身计为 FAIL；预期失败的调用传 ok=False。"""
    r = subprocess.run([PY, *map(str, a)], capture_output=True, text=True,
                       encoding="utf-8", errors="replace", env=env)
    if ok and r.returncode != 0:
        fails.append(f"exit!=0: {Path(str(a[0])).name} rc={r.returncode}")
        print("    ! cmd nonzero:", [str(x) for x in a], "\n   ",
              (r.stdout or "")[:200], (r.stderr or "")[:200])
    return r

def report(files, modified_at=None):
    return {"files": [{"path": p, "name": Path(p).name, "extension": Path(p).suffix,
                       "category": c, "size_bytes": 0,
                       **({"modified_at": modified_at} if modified_at else {})}
                      for p, c in files]}

def jload(p):
    return json.loads(Path(p).read_text(encoding="utf-8"))

def jload_str(s):
    return json.loads(s)


def main():
    W = Path(tempfile.mkdtemp(prefix="lmw_test_"))
    try:
        with sect("[1] scan_folder"):
            s1 = W / "s1"; s1.mkdir()
            (s1 / "a.md").write_text("# A", encoding="utf-8")
            (s1 / "b.txt").write_text("文本", encoding="utf-8")
            (s1 / "c.json").write_text('{"k":"v"}', encoding="utf-8")
            (s1 / "d.csv").write_text("h1,h2\n1,2", encoding="utf-8")
            (s1 / "e.png").write_text("x", encoding="utf-8")
            run(SC / "scan_folder.py", s1, "-o", W / "s1.json")
            d = jload(W / "s1.json")
            check("supported=4", d["supported_files"] == 4, d.get("supported_files"))
            check("unsupported=1", d["unsupported_files"] == 1)
            check("csv 归类 csv", d["categories"]["csv"]["count"] == 1)
            check("plan single_batch(<=20)", d["plan"]["strategy"] == "single_batch")
            check("非增量报告无 orphaned 字段", "orphaned" not in d)

            s2 = W / "s2"; s2.mkdir()
            for i in range(25):
                (s2 / f"f{i}.md").write_text(f"d{i}", encoding="utf-8")
            run(SC / "scan_folder.py", s2, "-o", W / "s2.json")
            check("plan multi_batch(21-100)", jload(W / "s2.json")["plan"]["strategy"] == "multi_batch")

            s3 = W / "s3"; s3.mkdir()
            for i in range(105):
                (s3 / f"f{i}.md").write_text(f"d{i}", encoding="utf-8")
            run(SC / "scan_folder.py", s3, "-o", W / "s3.json")
            p3 = jload(W / "s3.json")
            check("plan priority(>100)", p3["plan"]["strategy"] == "priority_batch")
            check("deferred_count=5", p3["plan"].get("deferred_count") == 5, p3["plan"].get("deferred_count"))
            check("M2: files 截断到 100", len(p3["files"]) == 100, len(p3["files"]))

        with sect("[2] generate_wiki_structure"):
            v = W / "vault"; run(SC / "generate_wiki_structure.py", "--output", v)
            for dd in ("documents", "summaries", "entities", "relations", ".obsidian", ".wiki-tree/extracted"):
                check(f"目录 {dd}", (v / dd).is_dir())
            check("空 manifest", jload(v / ".wiki-tree/manifest.json")["processed"] == {})
            idx = v / "_index.md"; idx.write_text(idx.read_text(encoding="utf-8") + "\nKEEP\n", encoding="utf-8")
            run(SC / "generate_wiki_structure.py", "--output", v)
            check("create-if-missing 不覆盖", "KEEP" in idx.read_text(encoding="utf-8"))
            run(SC / "generate_wiki_structure.py", "--output", v, "--force")
            check("--force 覆盖", "KEEP" not in idx.read_text(encoding="utf-8"))
            check("F2: 模板用 [[<kind>- 而非 [[entity-",
                  "[[entity-" not in (v / "relations/_knowledge-graph.md").read_text(encoding="utf-8"))
            gj = jload(v / ".obsidian/graph.json")
            qs = {g["query"] for g in gj["colorGroups"]}
            check("L7: 图谱含 date/location/event 配色",
                  {"tag:#date", "tag:#location", "tag:#event"} <= qs)
            check("低4: graph.json 用 Obsidian 原生 colorGroups 键、无 kebab 旧键",
                  "color-groups" not in gj)
            check("低4: 不生成数组形态 types.json", not (v / ".obsidian/types.json").exists())
            # 迁移：旧数组形态 types.json 重跑被清理；对象形态（Obsidian 自己维护）保留
            (v / ".obsidian/types.json").write_text('{"types": [{"name": "X"}]}', encoding="utf-8")
            run(SC / "generate_wiki_structure.py", "--output", v)
            check("低4: 旧数组 types.json 重跑被清理", not (v / ".obsidian/types.json").exists())
            (v / ".obsidian/types.json").write_text('{"types": {"created_at": "datetime"}}', encoding="utf-8")
            run(SC / "generate_wiki_structure.py", "--output", v)
            check("低4: 对象形态 types.json 保留", (v / ".obsidian/types.json").exists())

        with sect("[3] convert_documents"):
            cs = W / "cs"; cs.mkdir()
            (cs / "m.md").write_text("# M\n内容", encoding="utf-8")
            (cs / "t.csv").write_text("姓名,备注\n张三,负责 a|b\n", encoding="utf-8")
            (cs / "dupA.txt").write_text("相同内容", encoding="utf-8")
            (cs / "dupB.txt").write_text("相同内容", encoding="utf-8")
            (cs / "gbk.md").write_bytes("# 标题\nGBK中文内容".encode("gbk"))   # L2
            (cs / "fake.docx").write_text("not a real docx", encoding="utf-8")  # M1
            rep = W / "cs.json"
            rep.write_text(json.dumps(report([
                (str(cs / "m.md"), "markdown"), (str(cs / "t.csv"), "csv"),
                (str(cs / "dupA.txt"), "text"), (str(cs / "dupB.txt"), "text"),
                (str(cs / "gbk.md"), "markdown"), (str(cs / "fake.docx"), "word"),
            ]), ensure_ascii=False), encoding="utf-8")
            cv = W / "cv"; run(SC / "convert_documents.py", "--scan-report", rep, "--output", cv)
            cr = jload(cv / "_conversion_report.json")
            det = {Path(x["source"]).name: x for x in cr["details"]}
            check("csv→表格+竖线转义",
                  "| 姓名 | 备注 |" in (cv / "documents/t.md").read_text(encoding="utf-8")
                  and "a\\|b" in (cv / "documents/t.md").read_text(encoding="utf-8"))
            check("dupB 内容重复被跳过", det["dupB.txt"]["status"] == "skipped")
            check("L2: GBK 编码 .md 转换成功", det["gbk.md"]["status"] == "success"
                  and "GBK中文内容" in (cv / "documents/gbk.md").read_text(encoding="utf-8"))
            check("M1: 缺依赖/坏 .docx → error（非 success）", det["fake.docx"]["status"] == "error", det["fake.docx"]["status"])
            check("M1: 失败文档未写出 .md", not (cv / "documents/fake.md").exists())
            # 同源重转不产生 -1
            run(SC / "convert_documents.py", "--scan-report", rep, "--output", cv)
            check("同源重转无 -1 重复", not list((cv / "documents").glob("*-1.md")))

        with sect("[3b] convert 真实 docx (L8)"):
            try:
                import docx as _docx  # python-docx
                _have_docx = True
            except ImportError:
                _have_docx = False
            if not _have_docx:
                skip("真实 docx 测试（python-docx 未安装）")
            else:
                dd = W / "dd"; dd.mkdir()
                docp = dd / "real.docx"
                _d = _docx.Document()
                _d.add_heading("报告标题", level=1)
                _d.add_paragraph("正文段落内容。")
                _t = _d.add_table(rows=2, cols=2)
                _t.rows[0].cells[0].text = "姓名"; _t.rows[0].cells[1].text = "年龄"
                _t.rows[1].cells[0].text = "张三"; _t.rows[1].cells[1].text = "30"
                _d.add_table(rows=2, cols=2)  # 全空表 → 应被跳过且不占序号
                _t3 = _d.add_table(rows=2, cols=2)
                _t3.rows[0].cells[0].text = "项目"; _t3.rows[0].cells[1].text = "值"
                _t3.rows[1].cells[0].text = "备注"; _t3.rows[1].cells[1].text = "x|y"
                _d.save(str(docp))
                repd = W / "dd.json"
                repd.write_text(json.dumps(report([(str(docp), "word")]), ensure_ascii=False), encoding="utf-8")
                dv = W / "dv"; run(SC / "convert_documents.py", "--scan-report", repd, "--output", dv)
                crd = jload(dv / "_conversion_report.json")
                st = crd["details"][0]["status"]
                md = (dv / "documents/real.md").read_text(encoding="utf-8") if (dv / "documents/real.md").exists() else ""
                check("M1: 有效 docx 转换成功", st == "success", st)
                check("docx 段落保留", "正文段落内容" in md)
                check("L8: 表格渲染为 Markdown 表格", "| 姓名 | 年龄 |" in md and "| 张三 | 30 |" in md, md[:300])
                check("L8: 单元格竖线转义", "x\\|y" in md, md[:300])
                check("L8: 空表跳过 + 序号连续", "<!-- 表格 1 -->" in md and "<!-- 表格 2 -->" in md and "<!-- 表格 3 -->" not in md, md[:400])

        with sect("[4] update_manifest"):
            mv = W / "mv"; run(SC / "generate_wiki_structure.py", "--output", mv)
            mf = W / "mf.json"
            mf.write_text(json.dumps([
                {"source_path": str(cs / "m.md"), "doc_md": "documents/m.md", "doc_id": "m"},
                {"source_path": str(cs / "t.csv"), "doc_md": "documents/t.md", "doc_id": "t"},
            ], ensure_ascii=False), encoding="utf-8")
            run(SC / "update_manifest.py", "--vault", mv, "--mark-from", mf)
            man = jload(mv / ".wiki-tree/manifest.json")
            check("mark-from 登记 2 且 status=done",
                  len(man["processed"]) == 2 and all(x["status"] == "done" for x in man["processed"].values()))

            # L4: 损坏 manifest → 非破坏性（带时间戳备份，连续损坏不互相覆盖）
            bad = W / "badv"; (bad / ".wiki-tree").mkdir(parents=True)
            (bad / ".wiki-tree/manifest.json").write_text("{ 这不是合法 json", encoding="utf-8")
            run(SC / "update_manifest.py", "--vault", bad, "--mark", str(cs / "m.md"), "--doc-md", "documents/m.md")
            bk = list((bad / ".wiki-tree").glob("manifest.corrupt-*.json"))
            check("L4: 损坏 manifest 已备份(非破坏)", len(bk) == 1, [b.name for b in bk])
            check("L4: 备份保留损坏内容", bool(bk) and "这不是合法 json" in bk[0].read_text(encoding="utf-8"))
            check("L4: 重置后新 mark 写入成功", len(jload(bad / ".wiki-tree/manifest.json")["processed"]) == 1)
            (bad / ".wiki-tree/manifest.json").write_text("{ 再次损坏", encoding="utf-8")
            run(SC / "update_manifest.py", "--vault", bad, "--mark", str(cs / "dupA.txt"), "--doc-md", "documents/dupA.md")
            check("L4: 连续损坏产生 2 个独立备份",
                  len(list((bad / ".wiki-tree").glob("manifest.corrupt-*.json"))) == 2)

        with sect("[5] verify_entities"):
            doc = W / "ve.md"
            doc.write_text("---\nx: 1\n---\n\n项目用 C++ 实现。我们 WAIT 数据 available。上下文工程很重要。", encoding="utf-8")
            ents = W / "ve.json"
            ents.write_text(json.dumps({"entities": [
                {"kind": "concept", "text": "C++"},      # 符号实体、确实出现 → 应保留(L6)
                {"kind": "concept", "text": "AI"},       # 仅 WAIT/available 子串 → 应丢
                {"kind": "concept", "text": "上下文"},    # 中文子串 → 保留
            ]}, ensure_ascii=False), encoding="utf-8")
            res = jload_str(run(SC / "verify_entities.py", "--doc", doc, "--entities", ents).stdout)
            kept = {e["text"] for e in res["entities"]}
            check("L6: C++ 符号实体被保留", "C++" in kept, str(kept))
            check("AI 子串误放被丢", "AI" not in kept)
            check("上下文 中文子串保留", "上下文" in kept)
            # L6 已知取舍（有意行为，固化以防回归）：粘连更多 ASCII 词字符时不命中
            doc2 = W / "ve2.md"; doc2.write_text("用 C++11 和 ASP.NET 开发。", encoding="utf-8")
            ents2 = W / "ve2.json"
            ents2.write_text(json.dumps({"entities": [{"kind": "concept", "text": "C++"},
                                                       {"kind": "concept", "text": ".NET"}]}, ensure_ascii=False), encoding="utf-8")
            kept2 = {e["text"] for e in jload_str(run(SC / "verify_entities.py", "--doc", doc2, "--entities", ents2).stdout)["entities"]}
            check("L6 取舍(已知): C++ 不命中 C++11、.NET 不命中 ASP.NET", "C++" not in kept2 and ".NET" not in kept2, str(kept2))
            # H3 修复：边界限定 ASCII 词字符 → 紧贴 CJK 的真实实体不再被误杀
            doc3 = W / "ve3.md"; doc3.write_text("AI的应用很广泛，基于RAG的系统已部署。", encoding="utf-8")
            ents3 = W / "ve3.json"
            ents3.write_text(json.dumps({"entities": [{"kind": "concept", "text": "AI"},
                                                       {"kind": "concept", "text": "RAG"}]}, ensure_ascii=False), encoding="utf-8")
            kept3 = {e["text"] for e in jload_str(run(SC / "verify_entities.py", "--doc", doc3, "--entities", ents3).stdout)["entities"]}
            check("H3修复: AI/RAG 紧贴汉字命中（中文混排不再误杀）",
                  "AI" in kept3 and "RAG" in kept3, str(kept3))

        with sect("[6] compute_centrality"):
            cev = W / "cev"; (cev / ".wiki-tree/extracted").mkdir(parents=True)
            (cev / ".wiki-tree/extracted/d1.json").write_text(json.dumps({"doc_id": "d1",
                "entities": [{"kind": "concept", "text": t} for t in ("A", "B", "C")],
                "relations": [{"subject": "A", "predicate": "USES", "object": "B"},
                              {"subject": "A", "predicate": "USES", "object": "C"}]}, ensure_ascii=False), encoding="utf-8")
            (cev / ".wiki-tree/extracted/d2.json").write_text(json.dumps({"doc_id": "d2",
                "entities": [{"kind": "concept", "text": t} for t in ("A", "B")],
                "relations": [{"subject": "A", "predicate": "USES", "object": "B"}]}, ensure_ascii=False), encoding="utf-8")
            run(SC / "compute_centrality.py", "--vault", cev, "-o", W / "cen.json")
            co = jload(W / "cen.json")
            A = next(r for r in co["top"] if r["entity"] == "A")
            check("F4: A.degree=2(不同邻居数)", A["degree"] == 2, A["degree"])
            check("F4: A.relation_count=3(边条数)", A["relation_count"] == 3, A["relation_count"])
            check("degree != relation_count", A["degree"] != A["relation_count"])
            check("A 排第一", co["top"][0]["entity"] == "A")

        with sect("[7] F1 跨运行去重"):
            f1 = W / "f1"; (f1 / "sub").mkdir(parents=True)
            (f1 / "doc.md").write_text("唯一X", encoding="utf-8")
            (f1 / "sub" / "copy.md").write_text("唯一X", encoding="utf-8")
            (f1 / "uniq.md").write_text("内容Y", encoding="utf-8")
            fv = W / "fv"; run(SC / "generate_wiki_structure.py", "--output", fv)
            run(SC / "scan_folder.py", f1, "-o", fv / ".wiki-tree/scan.json")
            run(SC / "convert_documents.py", "--scan-report", fv / ".wiki-tree/scan.json", "--output", fv)
            rc = jload(fv / "_conversion_report.json")
            dup = next((x for x in rc["details"] if x["status"] == "skipped"), None)
            check("copy.md 被判内容副本", dup is not None)
            canon = "documents/" + Path(dup["duplicate_of"]).name
            marks = [{"source_path": str(f1 / "doc.md"), "doc_md": "documents/doc.md", "doc_id": "doc"},
                     {"source_path": str(f1 / "uniq.md"), "doc_md": "documents/uniq.md", "doc_id": "uniq"},
                     {"source_path": dup["source"], "doc_md": canon, "doc_id": Path(canon).stem}]
            (fv / ".wiki-tree/_marks.json").write_text(json.dumps(marks, ensure_ascii=False), encoding="utf-8")
            run(SC / "update_manifest.py", "--vault", fv, "--mark-from", fv / ".wiki-tree/_marks.json")
            run(SC / "scan_folder.py", f1, "--vault", fv, "-o", fv / ".wiki-tree/scan2.json")
            check("F1: 副本登记后第二轮 pending=0", jload(fv / ".wiki-tree/scan2.json")["pending_count"] == 0)

        with sect("[8] F3 UTF-8 stdout"):
            clean = {k: vv for k, vv in os.environ.items() if k not in ("PYTHONUTF8", "PYTHONIOENCODING")}
            cn = W / "cn"; cn.mkdir(); (cn / "照明数据.md").write_text("x", encoding="utf-8")
            r = run(SC / "scan_folder.py", cn, env=clean)
            check("scan stdout 中文文件名无乱码", "照明数据" in r.stdout and r.returncode == 0, (r.stdout or "")[:60])

        with sect("[9] 补充覆盖"):
            # 9a convert_json 递归渲染
            cj = W / "cj"; cj.mkdir()
            (cj / "data.json").write_text(json.dumps({"name": "张三", "meta": {"role": "工程师"}}, ensure_ascii=False), encoding="utf-8")
            repj = W / "cj.json"
            repj.write_text(json.dumps(report([(str(cj / "data.json"), "json")]), ensure_ascii=False), encoding="utf-8")
            jv = W / "jv"; run(SC / "convert_documents.py", "--scan-report", repj, "--output", jv)
            jmd = (jv / "documents/data.md").read_text(encoding="utf-8")
            check("convert_json: 键值递归渲染", "**name**: 张三" in jmd and "role" in jmd, jmd[:200])

            # 9b CSV 超大表截断 (CSV_ROW_CAP=1000)
            big = W / "big"; big.mkdir()
            (big / "big.csv").write_text("\n".join(["col"] + [f"ROW{i:04d}" for i in range(1, 1101)]), encoding="utf-8")
            repb = W / "big.json"
            repb.write_text(json.dumps(report([(str(big / "big.csv"), "csv")]), ensure_ascii=False), encoding="utf-8")
            bv = W / "bv"; run(SC / "convert_documents.py", "--scan-report", repb, "--output", bv)
            bmd = (bv / "documents/big.md").read_text(encoding="utf-8")
            check("CSV 截断: 含过大提示(共1100)", "表格过大" in bmd and "1100" in bmd)
            check("CSV 截断: 保留 ROW0001、丢弃 ROW1100", "ROW0001" in bmd and "ROW1100" not in bmd)

            # 9c convert_pdf（PyMuPDF 可用时实测，否则跳过）
            try:
                import fitz as _fitz
                _have_pdf = True
            except ImportError:
                _have_pdf = False
            if not _have_pdf:
                skip("convert_pdf 测试（PyMuPDF 未安装）")
            else:
                pd = W / "pd"; pd.mkdir(); pdfp = pd / "doc.pdf"
                _pdoc = _fitz.open()
                _pdoc.new_page().insert_text((72, 72), "PDF page one text")
                _pdoc.save(str(pdfp)); _pdoc.close()
                # 无文字层（扫描版形态）与加密 PDF
                blankp = pd / "blank.pdf"
                _pb = _fitz.open(); _pb.new_page(); _pb.save(str(blankp)); _pb.close()
                encp = pd / "enc.pdf"
                _pe = _fitz.open(); _pe.new_page().insert_text((72, 72), "secret")
                _pe.save(str(encp), encryption=_fitz.PDF_ENCRYPT_AES_256, owner_pw="o", user_pw="u"); _pe.close()
                repp = W / "pd.json"
                repp.write_text(json.dumps(report([(str(pdfp), "pdf"), (str(blankp), "pdf"),
                                                   (str(encp), "pdf")]), ensure_ascii=False), encoding="utf-8")
                pv = W / "pv"; run(SC / "convert_documents.py", "--scan-report", repp, "--output", pv)
                pmd = (pv / "documents/doc.md").read_text(encoding="utf-8") if (pv / "documents/doc.md").exists() else ""
                check("convert_pdf: 页码标注 + 文本", "<!-- 第 1 页 -->" in pmd and "PDF page one text" in pmd, pmd[:200])
                pdet = {Path(x["source"]).name: x for x in jload(pv / "_conversion_report.json")["details"]}
                check("中11: 无文字层 PDF → empty（不写盘不标 success）",
                      pdet["blank.pdf"]["status"] == "empty" and not (pv / "documents/blank.md").exists(),
                      pdet["blank.pdf"])
                check("中11: 加密 PDF → error 且报已加密",
                      pdet["enc.pdf"]["status"] == "error" and "已加密" in pdet["enc.pdf"]["reason"], pdet["enc.pdf"])

            # 9d 增量 modified：源 mtime 变化 → 重新待处理
            ms = W / "ms"; ms.mkdir(); (ms / "x.md").write_text("内容X", encoding="utf-8")
            mvt = W / "mvt"; run(SC / "generate_wiki_structure.py", "--output", mvt)
            run(SC / "scan_folder.py", ms, "-o", mvt / ".wiki-tree/scan.json")
            run(SC / "update_manifest.py", "--vault", mvt, "--mark", str(ms / "x.md"), "--doc-md", "documents/x.md")
            run(SC / "scan_folder.py", ms, "--vault", mvt, "-o", mvt / ".wiki-tree/s1.json")
            inc1 = jload(mvt / ".wiki-tree/s1.json")
            check("增量: 未改动 → done(pending 0) 且 orphaned 为空",
                  inc1["pending_count"] == 0 and inc1["orphaned"] == [] and inc1["orphaned_count"] == 0)
            os.utime(ms / "x.md", (1_700_000_000, 1_700_000_000))  # 强制改 mtime
            run(SC / "scan_folder.py", ms, "--vault", mvt, "-o", mvt / ".wiki-tree/s2.json")
            s2r = jload(mvt / ".wiki-tree/s2.json")
            st = {Path(f["path"]).name: f["status"] for f in s2r["files"]}
            check("增量: mtime 变化 → modified(重新 pending)", st.get("x.md") == "modified" and s2r["pending_count"] == 1, str(st))

            # 9e compute_centrality --dedup-map：变体合并
            dmp = W / "dmap.json"; dmp.write_text(json.dumps({"C": "B"}, ensure_ascii=False), encoding="utf-8")
            run(SC / "compute_centrality.py", "--vault", cev, "--dedup-map", dmp, "-o", W / "cen_dm.json")
            co_dm = jload(W / "cen_dm.json")
            ents_dm = {r["entity"] for r in co_dm["top"]}
            check("dedup-map: C 合并入 B（C 消失、dedup_applied）", "C" not in ents_dm and co_dm["dedup_applied"] is True, str(ents_dm))

        with sect("[10] scan exclude/.mwignore"):
            ex = W / "ex"; (ex / "nested").mkdir(parents=True); (ex / "sub").mkdir()
            (ex / "a.md").write_text("# A", encoding="utf-8")
            (ex / "sub" / "b.md").write_text("# B", encoding="utf-8")
            (ex / "nested" / "skill.md").write_text("# nested", encoding="utf-8")
            run(SC / "scan_folder.py", ex, "--exclude", "nested", "-o", W / "ex1.json")
            names1 = {f["name"] for f in jload(W / "ex1.json")["files"]}
            check("exclude 目录: nested/ 被剪枝", "skill.md" not in names1 and {"a.md", "b.md"} <= names1, str(names1))
            run(SC / "scan_folder.py", ex, "--exclude", "skill.md", "-o", W / "ex2.json")
            e2 = jload(W / "ex2.json")
            check("exclude 文件名: skill.md 跳过 + excluded_count>=1",
                  "skill.md" not in {f["name"] for f in e2["files"]} and e2["excluded_count"] >= 1, e2.get("excluded_count"))
            (ex / ".mwignore").write_text("# 注释\nnested/\n", encoding="utf-8")
            run(SC / "scan_folder.py", ex, "-o", W / "ex3.json")
            check(".mwignore: nested/ 被排除", "skill.md" not in {f["name"] for f in jload(W / "ex3.json")["files"]})

        with sect("[11] convert front-matter 降级"):
            fmd = W / "fmd"; fmd.mkdir()
            (fmd / "withfm.md").write_text("---\nname: lighting-x\ntags:\n  - t1\n---\n\n# 标题\n正文 lighting-x 提及。", encoding="utf-8")
            (fmd / "plain.md").write_text("# 无fm\n正文", encoding="utf-8")
            repfm = W / "fmd.json"
            repfm.write_text(json.dumps(report([(str(fmd / "withfm.md"), "markdown"), (str(fmd / "plain.md"), "markdown")]), ensure_ascii=False), encoding="utf-8")
            fmv = W / "fmv"; run(SC / "convert_documents.py", "--scan-report", repfm, "--output", fmv)
            wmd = (fmv / "documents/withfm.md").read_text(encoding="utf-8")
            ndash = sum(1 for ln in wmd.splitlines() if ln.strip() == "---")
            check("front-matter 降级: 全文只剩一对 YAML 头（无双 fm）", ndash == 2, ndash)
            check("front-matter 降级: 源 fm 降级为引用块", "> name: lighting-x" in wmd)
            check("front-matter 降级: 源 fm 不再是裸 YAML 行", "\nname: lighting-x" not in wmd)
            check("front-matter 降级: 实体仍在正文可匹配", "lighting-x 提及" in wmd)
            check("无 fm 的 .md 不受影响", "# 无fm" in (fmv / "documents/plain.md").read_text(encoding="utf-8"))

        with sect("[12] update_manifest --from-conversion-report"):
            crv = W / "crv"; run(SC / "generate_wiki_structure.py", "--output", crv)
            fake_cr = {"details": [
                {"status": "success", "source": str(cs / "m.md"), "output": str(crv / "documents/m.md")},
                {"status": "skipped", "source": str(cs / "dupB.txt"), "duplicate_of": str(crv / "documents/dupA.md")},
                {"status": "error", "source": str(cs / "fake.docx"), "reason": "x"},
            ]}
            crj = W / "fake_cr.json"; crj.write_text(json.dumps(fake_cr, ensure_ascii=False), encoding="utf-8")
            run(SC / "update_manifest.py", "--vault", crv, "--from-conversion-report", crj)
            cman = jload(crv / ".wiki-tree/manifest.json")["processed"]
            check("from-conversion-report: success+去重副本各登记(2)、error 不登记", len(cman) == 2, len(cman))
            dup_rec = cman.get(str(Path(str(cs / "dupB.txt")).resolve()))
            check("from-conversion-report: 副本 doc_md 指向 canonical",
                  bool(dup_rec) and dup_rec["doc_md"] == "documents/dupA.md", dup_rec)
            check("from-conversion-report: 无 modified_at 时回退 stat（mtime 非空）",
                  all(x.get("mtime") for x in cman.values()), cman)
            cmv = W / "cmv"; run(SC / "generate_wiki_structure.py", "--output", cmv)
            marks_f = W / "marks_clean.json"
            marks_f.write_text(json.dumps([{"source_path": str(cs / "m.md"), "doc_md": "documents/m.md", "doc_id": "m"}], ensure_ascii=False), encoding="utf-8")
            run(SC / "update_manifest.py", "--vault", cmv, "--mark-from", marks_f, "--clean-marks")
            check("--clean-marks: 清单文件被删除", not marks_f.exists())
            check("--clean-marks: 仍正确登记", len(jload(cmv / ".wiki-tree/manifest.json")["processed"]) == 1)

        with sect("[13] assemble_vault"):
            av = W / "av"; run(SC / "generate_wiki_structure.py", "--output", av)
            (av / "documents/doc1.md").write_text("# doc1\n项目P 工具T 概念C", encoding="utf-8")
            (av / "documents/doc2.md").write_text("# doc2\n工具T 概念C", encoding="utf-8")
            exd = av / ".wiki-tree/extracted"
            (exd / "doc1.json").write_text(json.dumps({
                "doc_id": "doc1", "doc_md": "documents/doc1.md", "short_summary": "文档一概要",
                "detailed_summary": "文档一详细摘要：项目P 使用 工具T 与 概念C。", "importance": 0.9,
                "topics": ["主题甲"],
                "entities": [{"kind": "project", "text": "项目P"}, {"kind": "tool", "text": "工具T"}, {"kind": "concept", "text": "概念C"}],
                "relations": [{"subject": "项目P", "predicate": "USES", "object": "工具T", "confidence": 0.9, "evidence": "P 用 T"},
                              {"subject": "项目P", "predicate": "USES", "object": "概念C", "confidence": 0.8}]}, ensure_ascii=False), encoding="utf-8")
            (exd / "doc2.json").write_text(json.dumps({
                "doc_id": "doc2", "doc_md": "documents/doc2.md", "short_summary": "文档二概要",
                "detailed_summary": "文档二详细摘要：工具T 关联 概念C。", "importance": 0.5,
                "topics": ["主题甲"],
                "entities": [{"kind": "tool", "text": "工具T"}, {"kind": "concept", "text": "概念C"}],
                "relations": [{"subject": "工具T", "predicate": "RELATED_TO", "object": "概念C", "confidence": 0.7}]}, ensure_ascii=False), encoding="utf-8")
            asum = jload_str(run(SC / "assemble_vault.py", "--vault", av).stdout)
            check("assemble: 文档/实体/关系计数", asum["docs"] == 2 and asum["entities"] == 3 and asum["relations"] == 3, asum)
            check("assemble: 生成 index/graph/report",
                  (av / "_index.md").exists() and (av / "relations/_knowledge-graph.md").exists() and (av / "_processing-report.md").exists())
            check("assemble: _index 统计回填", "文档总数**：2" in (av / "_index.md").read_text(encoding="utf-8"))
            check("assemble: 卡片建给连通实体(degree>=1)", asum["cards_created"] >= 3, asum["cards_created"])
            notes = {p.stem for p in av.rglob("*.md") if ".wiki-tree" not in p.parts}
            dangling = []
            for p in av.rglob("*.md"):
                if ".wiki-tree" in p.parts:
                    continue
                for m in re.finditer(r"\[\[([^\]]+)\]\]", p.read_text(encoding="utf-8")):
                    tgt = m.group(1).split("|")[0].split("#")[0].strip().split("/")[-1]
                    if tgt not in notes:
                        dangling.append((p.name, tgt))
            check("assemble: 零悬空 wikilink", not dangling, str(dangling[:5]))
            ra2 = jload_str(run(SC / "assemble_vault.py", "--vault", av).stdout)
            check("中14: 再跑受管刷新（updated=3、不重建不跳过）",
                  ra2["cards_created"] == 0 and ra2["cards_updated"] >= 3 and ra2["cards_skipped"] == 0, ra2)
            rep_txt = (av / "_processing-report.md").read_text(encoding="utf-8")
            check("assemble: importance 被消费（doc1=0.9 排在 doc2=0.5 前）", rep_txt.index("doc1") < rep_txt.index("doc2"))
            # 中14: 标记外散文保留 + 受管刷新幂等
            card = av / "entities/tool-工具T.md"
            card.write_text(card.read_text(encoding="utf-8") + "\n手写散文段。\n", encoding="utf-8")
            run(SC / "assemble_vault.py", "--vault", av)
            c3 = card.read_text(encoding="utf-8")
            check("中14: 标记外散文保留", "手写散文段。" in c3 and c3.count("<!-- wiki-tree:auto:start -->") == 1)
            run(SC / "assemble_vault.py", "--vault", av)
            check("中14: 受管刷新字节级幂等", card.read_text(encoding="utf-8") == c3)
            # 中14: 无标记旧卡兼容（不动内容、计 skipped）
            legacy_card = av / "entities/concept-概念C.md"
            legacy_card.write_text("# 概念C\n纯手工内容\n", encoding="utf-8")
            ra5 = jload_str(run(SC / "assemble_vault.py", "--vault", av).stdout)
            check("中14: 无标记旧卡跳过不覆盖",
                  legacy_card.read_text(encoding="utf-8") == "# 概念C\n纯手工内容\n" and ra5["cards_skipped"] == 1, ra5)

        with sect("[14] suggest_dedup"):
            sv = W / "sv"; (sv / ".wiki-tree/extracted").mkdir(parents=True)
            (sv / ".wiki-tree/extracted/d.json").write_text(json.dumps({"doc_id": "d",
                "entities": [{"kind": "tool", "text": "Chart.js"}, {"kind": "tool", "text": "chartjs"},
                             {"kind": "concept", "text": "RAG"}, {"kind": "concept", "text": "RAG系统"},
                             {"kind": "tool", "text": "Python"}, {"kind": "tool", "text": "Pyton"},
                             {"kind": "tool", "text": 123}]}, ensure_ascii=False), encoding="utf-8")
            sd = jload_str(run(SC / "suggest_dedup.py", "--vault", sv).stdout)
            cand = {tuple(sorted(c["variants"])): c.get("reason", "") for c in sd["candidates"]}
            check("suggest_dedup: 归一化相等候选(Chart.js/chartjs)",
                  "归一化" in cand.get(tuple(sorted(("Chart.js", "chartjs"))), ""), str(cand))
            check("suggest_dedup: 子串候选(RAG/RAG系统)",
                  any("RAG" in p and "RAG系统" in p and "子串" in r for p, r in cand.items()), str(cand))
            check("低8: 编辑距离候选(Python/Pyton) 规则不漂移",
                  any(set(p) == {"Python", "Pyton"} and r.startswith("编辑距离") for p, r in cand.items()), str(cand))

        with sect("[15] emit_access_bundle + kb_query"):
            # 前置：emit 硬依赖 centrality.json 与 summaries/topic-*.md（Phase 5/6 产物），复用 [13] 的 av vault
            run(SC / "compute_centrality.py", "--vault", av, "-o", av / ".wiki-tree/centrality.json")
            (av / "summaries/topic-主题甲.md").write_text(
                "---\nkind: summary\ntopic: 主题甲\ndoc_count: 2\n---\n\n"
                "**一句话**：围绕项目P与工具T的测试主题。\n\n要点：项目P 使用 工具T。\n", encoding="utf-8")
            (av / "summaries/topic-主题丙.md").write_text(
                "# 主题丙\n\n这是回退摘要。第二句不该出现。\n", encoding="utf-8")
            shutil.copy(SC / "templates/kb_query.py", av / "kb_query.py")
            shutil.copy(SC / "templates/kb_mcp_server.py", av / "kb_mcp_server.py")
            run(SC / "emit_access_bundle.py", "--vault", av, "--id", "t", "--name", "测试库",
                "--scope", "用于测试的知识库", "--extra-use-when", "额外词")
            kbj = jload(av / "kb.json")
            check("emit: stats(2 文档/3 实体/3 关系)",
                  kbj["stats"]["documents"] == 2 and kbj["stats"]["entities"] == 3 and kbj["stats"]["relations"] == 3,
                  kbj["stats"])
            t0 = next((t for t in kbj["topics"] if t["name"] == "主题甲"), None)
            check("emit: topic 主题甲(docs=2) + one_liner 解析",
                  bool(t0) and t0["docs"] == 2 and t0["one_liner"] == "围绕项目P与工具T的测试主题。", t0)
            t3 = next((t for t in kbj["topics"] if t["name"] == "主题丙"), None)
            check("中8: one_liner 缺「**一句话**」行 → 回退首段首句",
                  bool(t3) and t3["one_liner"] == "这是回退摘要", t3)
            check("emit: use_when=中心度实体+主题+extra",
                  {"项目P", "工具T", "概念C", "主题甲", "额外词"} <= set(kbj["use_when"]), kbj["use_when"])
            check("低13: extra_use_when 原始留档", kbj.get("extra_use_when") == ["额外词"], kbj.get("extra_use_when"))
            idxj = jload(av / ".wiki-tree/search-index.json")
            d1idx = next((r for r in idxj["docs"] if r["id"] == "doc1"), None)
            check("高4: search-index v2（bigram+df+停用字过滤+doc_id 进 blob）",
                  idxj.get("version") == 2 and idxj["df"].get("工具") == 2 and bool(d1idx)
                  and {"概念", "工具", "概", "p", "doc1"} <= set(d1idx["tok"]) and "与" not in set(d1idx["tok"]),
                  d1idx and d1idx["tok"][:10])
            check("emit: AGENTS.md/.mcp.json 生成",
                  "主题甲" in (av / "AGENTS.md").read_text(encoding="utf-8")
                  and "t-kb" in jload(av / ".mcp.json")["mcpServers"])
            check("低20: .mcp.json command=当前解释器",
                  jload(av / ".mcp.json")["mcpServers"]["t-kb"]["command"] == Path(sys.executable).as_posix())
            # kb_query 检索 + 四档下钻（kb.json 是 emit↔query 的隐式契约，这里两端对测）
            q15 = jload_str(run(av / "kb_query.py", "工具T 概念", "--json").stdout)
            check("query: 候选文档命中且 doc1(imp 0.9)排第一",
                  bool(q15["documents"]) and q15["documents"][0]["path"] == "documents/doc1.md",
                  [d["path"] for d in q15["documents"]])
            check("query: 主题层命中 主题甲", any(t["name"] == "主题甲" for t in q15["topics"]), q15["topics"])
            q15d = jload_str(run(av / "kb_query.py", "工具T 概念", "--json", "--level", "detailed").stdout)
            check("query: --level detailed 附逐文档详细摘要",
                  "详细摘要" in (q15d["documents"][0].get("detailed") or ""))
            check("query: --topic 取 L1 摘要", "一句话" in run(av / "kb_query.py", "--topic", "主题甲").stdout)
            check("高4: --topic 乱序单字「甲题」不误命中",
                  "未找到主题" in run(av / "kb_query.py", "--topic", "甲题").stdout)
            check("高4: --topic 前缀「主题」bigram 子集仍命中",
                  "一句话" in run(av / "kb_query.py", "--topic", "主题").stdout)
            ent_out = run(av / "kb_query.py", "--entity", "工具T").stdout
            check("query: --entity 取实体卡", "工具T" in ent_out and "未找到" not in ent_out, ent_out[:80])
            (av / "entities/00-[草稿]方案.md").write_text("# [草稿]方案\n", encoding="utf-8")
            eo = run(av / "kb_query.py", "--entity", "[草稿]方案").stdout
            check("低12: --entity 名含[ ]元字符可命中", "[草稿]方案" in eo and "未找到" not in eo, eo[:60])
            rg = run(av / "kb_query.py", "--global")
            check("query: --global 取 L2 全局摘要", rg.returncode == 0 and rg.stdout.strip() != "")
            gs = av / "summaries/_global-summary.md"
            gbak = gs.read_text(encoding="utf-8"); gs.unlink()
            rg2 = run(av / "kb_query.py", "--global")
            check("中21: --global 缺文件 → 友好提示非 traceback",
                  rg2.returncode == 0 and "未找到全局摘要" in rg2.stdout, (rg2.stdout or "")[:80])
            gs.write_text(gbak, encoding="utf-8")
            check("query: --list-topics", "主题甲" in run(av / "kb_query.py", "--list-topics").stdout)
            check("query: --doc full 取 L0 原文",
                  "项目P" in run(av / "kb_query.py", "--doc", "doc1", "--level", "full").stdout)
            # 高2: 索引过期 → stderr 警告 + 回退全扫
            time.sleep(0.05); os.utime(av / ".wiki-tree/extracted/doc1.json")
            rsq = run(av / "kb_query.py", "工具T 概念", "--json")
            check("高2: 索引过期 → stderr 警告 + 回退结果一致",
                  "索引过期" in (rsq.stderr or "") and jload_str(rsq.stdout)["documents"][0]["path"] == "documents/doc1.md",
                  (rsq.stderr or "")[:80])
            (av / ".wiki-tree/search-index.json").unlink()  # 索引缺失 → 应回退扫 extracted，结果一致
            q15f = jload_str(run(av / "kb_query.py", "工具T 概念", "--json").stdout)
            check("query: 无索引回退扫 extracted 结果一致",
                  bool(q15f["documents"]) and q15f["documents"][0]["path"] == "documents/doc1.md")
            # 中18: 增量引入新主题但无摘要文件 → summary_missing 提示而非悬空 404
            (av / ".wiki-tree/extracted/doc3.json").write_text(json.dumps({
                "doc_id": "doc3", "doc_md": "documents/doc3.md", "short_summary": "文档三概要",
                "detailed_summary": "文档三：主题乙相关。", "importance": 0.4, "topics": ["主题乙"],
                "entities": [], "relations": []}, ensure_ascii=False), encoding="utf-8")
            run(SC / "emit_access_bundle.py", "--vault", av, "--id", "t", "--name", "测试库", "--scope", "s")
            tb = next((t for t in jload(av / "kb.json")["topics"] if t["name"] == "主题乙"), None)
            check("中18: 无摘要主题标 summary_missing", bool(tb) and tb.get("summary_missing") is True, tb)
            check("中18: --topic 缺摘要 → 尚未生成提示", "尚未生成" in run(av / "kb_query.py", "--topic", "主题乙").stdout)
            qm = jload_str(run(av / "kb_query.py", "主题乙", "--json").stdout)
            check("中18: search 缺摘要主题 one_liner 提示",
                  any(t["name"] == "主题乙" and t["one_liner"] == "(摘要尚未生成)" for t in qm["topics"]), qm["topics"])
            (av / ".wiki-tree/extracted/doc3.json").unlink()
            # 还原 2 文档新鲜态（带 extra，供 [17] finalize 不动点测试消费）
            run(SC / "emit_access_bundle.py", "--vault", av, "--id", "t", "--name", "测试库",
                "--scope", "用于测试的知识库", "--extra-use-when", "额外词")
            # kb_mcp_server: 缺 MCP SDK 时应干净退出(rc=2)并指引 CLI 兜底（装了 SDK 则跳过；CI 裸腿覆盖此路径）
            try:
                import mcp as _mcp  # noqa: F401
                _have_mcp = True
            except ImportError:
                _have_mcp = False
            if _have_mcp:
                skip("kb_mcp_server 缺 SDK 退出路径（mcp 已安装）")
                # 中19/高4: kb_hub_server 进程内检索（与 kb_query 同打分；缓存键含 mtime+size）
                sys.path.insert(0, str(SC / "templates"))
                import kb_hub_server as hub
                hres = hub._search(str(av), "工具T 概念", 5, "short")
                check("hub: 检索 doc1 第一", bool(hres["documents"]) and hres["documents"][0]["path"] == "documents/doc1.md",
                      hres["documents"][:2])
                check("中19: 缓存键为 (mtime,size) 元组", isinstance(hub._IDX_CACHE[str(av)][0], tuple))
                check("hub: --topic 乱序不误命中", "未找到主题" in hub._topic(str(av), "甲题"))
            else:
                rm = run(av / "kb_mcp_server.py", ok=False)
                check("mcp_server: 缺 SDK → rc=2 + 提示 CLI 兜底",
                      rm.returncode == 2 and "kb_query" in rm.stderr, (rm.stderr or "")[:120])
                skip("kb_hub 进程内检索（mcp 未安装）")

        with sect("[16] kb_register（全程临时 registry/hook 文件，不触碰真实 ~）"):
            reg_p = W / "reg.json"; hook_p = W / "hook.md"
            hook_p.write_text("# 用户手写前置内容\n", encoding="utf-8")
            run(SC / "kb_register.py", "--vault", av, "--registry", reg_p, "--hook-file", hook_p, "--install-hook")
            run(SC / "kb_register.py", "--vault", av, "--registry", reg_p, "--hook-file", hook_p, "--install-hook")
            regd = jload(reg_p)
            check("register: 同 id 双跑 upsert 幂等(1 条)",
                  len(regd["knowledge_bases"]) == 1 and regd["knowledge_bases"][0]["id"] == "t",
                  [k["id"] for k in regd["knowledge_bases"]])
            check("register: 条目含 root + query_cli",
                  bool(regd["knowledge_bases"][0].get("root"))
                  and regd["knowledge_bases"][0]["query_cli"].endswith("--json"))
            ht = hook_p.read_text(encoding="utf-8")
            check("hook: 双跑后 BEGIN/END 各恰 1 个（无重复块）",
                  ht.count("<!-- KB-HUB:BEGIN") == 1 and ht.count("<!-- KB-HUB:END -->") == 1)
            check("hook: 块外用户内容保留", "用户手写前置内容" in ht)
            check("高7: 首跑产生备份、幂等重跑不再备份",
                  len(list(W.glob("hook.md.bak-*"))) == 1, [p.name for p in W.glob("hook.md.bak-*")])
            hook_p.write_text(ht + "\n# 用户手写后置内容\n", encoding="utf-8")
            vb = W / "vb"; vb.mkdir()
            (vb / "kb.json").write_text(json.dumps({
                "id": "b2", "name": "财务库", "scope": "发票报销与增值税财务知识",
                "use_when": ["发票", "报销", "增值税", "财务"], "stats": {}, "topics": []},
                ensure_ascii=False), encoding="utf-8")
            run(SC / "kb_register.py", "--vault", vb, "--registry", reg_p, "--hook-file", hook_p, "--install-hook")
            regd2 = jload(reg_p)
            check("register: 第二库登记 + 按 id 排序",
                  [k["id"] for k in regd2["knowledge_bases"]] == ["b2", "t"])
            ht2 = hook_p.read_text(encoding="utf-8")
            check("hook: 块内更新为 2 库、块外前/后置内容原样",
                  ht2.count("<!-- KB-HUB:BEGIN") == 1 and "**b2**" in ht2 and "**t**" in ht2
                  and "用户手写前置内容" in ht2 and "用户手写后置内容" in ht2)
            check("高7: 无 .tmp 残留（原子写收尾干净）",
                  not list(W.glob("*.tmp")) and not list(W.glob("hook.md.tmp*")))
            # 高7: registry 损坏 → fail loud，不静默清空
            reg_bad = W / "reg_bad.json"; reg_bad.write_text("{broken", encoding="utf-8")
            rb16 = run(SC / "kb_register.py", "--vault", av, "--registry", reg_bad,
                       "--hook-file", W / "h2.md", ok=False)
            check("高7: 损坏 registry → rc=1 + 原文件未被覆盖",
                  rb16.returncode == 1 and "JSON" in rb16.stderr
                  and reg_bad.read_text(encoding="utf-8") == "{broken", (rb16.stderr or "")[:100])
            # 高7: 孤立 BEGIN → 拒绝改写（追加会在下次替换时吞掉用户内容）
            hook_o = W / "hook_o.md"
            hook_o.write_text("# 前置\n<!-- KB-HUB:BEGIN auto-managed by wiki-tree -->\n用户被夹内容\n", encoding="utf-8")
            ro = run(SC / "kb_register.py", "--vault", av, "--registry", W / "reg_o.json",
                     "--hook-file", hook_o, "--install-hook", ok=False)
            check("高7: 孤立 BEGIN → rc=1 + 文件原样",
                  ro.returncode == 1 and "孤立" in ro.stderr
                  and "用户被夹内容" in hook_o.read_text(encoding="utf-8")
                  and "<!-- KB-HUB:END -->" not in hook_o.read_text(encoding="utf-8"), (ro.stderr or "")[:100])

        with sect("[17] kb_ingest（假 HOME，不触碰真实 ~/.knowledge-bases）"):
            KI = SC / "templates/kb_ingest.py"
            H = W / "home"; (H / ".knowledge-bases").mkdir(parents=True)
            henv = dict(os.environ, HOME=str(H), USERPROFILE=str(H))
            (H / ".knowledge-bases/ingest-policy.json").write_text(json.dumps({
                "mode": "confirm", "min_score": 2.0, "ratio": 1.5,
                "per_kb_override": {}, "skill_scripts_dir": str(SC)}, ensure_ascii=False), encoding="utf-8")
            hreg = H / ".knowledge-bases/registry.json"
            run(SC / "kb_register.py", "--vault", av, "--registry", hreg)
            run(SC / "kb_register.py", "--vault", vb, "--registry", hreg)
            src17 = W / "src17"; src17.mkdir()
            f_inv = src17 / "发票合规指南.md"
            f_inv.write_text("本文讲发票报销流程与增值税专用发票的合规要求。", encoding="utf-8")
            f_misc = src17 / "随笔.md"
            f_misc.write_text("今天天气晴朗，出门散步看云。", encoding="utf-8")
            pl = jload_str(run(KI, f_inv, f_misc, "--plan", "--json", env=henv).stdout)
            st17 = {p["name"]: p for p in pl["plan"]}
            check("ingest: 强信号文件高置信判入 b2",
                  st17["发票合规指南.md"]["state"] == "assign" and st17["发票合规指南.md"].get("target") == "b2",
                  st17["发票合规指南.md"])
            check("ingest: 无关文件判 无匹配(none·不写)",
                  st17["随笔.md"]["state"] == "none", st17["随笔.md"]["state"])
            sg = jload_str(run(KI, f_inv, "--kb", "b2", "--yes", "--json", env=henv).stdout)
            staged_p = vb / ".wiki-tree/_ingest/发票合规指南.md"
            check("ingest: --kb --yes 暂存进目标 inbox + 审计日志",
                  len(sg.get("staged", [])) == 1 and staged_p.exists()
                  and (H / ".knowledge-bases/ingest-log.jsonl").exists(), sg.get("staged"))
            check("ingest: 原始文件不受损", f_inv.exists() and "增值税" in f_inv.read_text(encoding="utf-8"))
            rb = run(KI, "--rollback", "发票合规指南.md", "--kb", "b2", env=henv)
            check("ingest: --rollback 删除暂存副本",
                  "已回滚" in rb.stdout and not staged_p.exists(), (rb.stdout or "")[:120])
            # 高6(a)+中20: --build 后 stdout 仍是纯 JSON；rollback 经报告映射删除产物（含空格文件名）
            f_sp = src17 / "会议纪要 2026.md"
            f_sp.write_text("会议讨论了下季度财务预算与报销流程。", encoding="utf-8")
            rb1 = run(KI, f_sp, "--kb", "b2", "--yes", "--build", "--json", env=henv)
            out_b = jload_str(rb1.stdout)  # stdout 必须整体是合法 JSON（中20）
            check("中20: --json --build stdout 纯 JSON + build_log",
                  out_b.get("built") is True and isinstance(out_b.get("build_log"), list)
                  and any("convert_documents.py" in l for l in out_b["build_log"]), str(out_b.get("built")))
            doc_sp = vb / "documents/会议纪要-2026.md"
            check("高6: build 产物已生成（空格→连字符）", doc_sp.exists())
            rb2 = run(KI, "--rollback", "会议纪要 2026.md", "--kb", "b2", env=henv)
            check("高6: rollback 经 details 映射删除产物（死代码复活）",
                  not doc_sp.exists() and "已回滚暂存" in rb2.stdout and "警告" not in (rb2.stderr or ""),
                  (rb2.stderr or "")[:120])
            # 高6(b): 报告缺失 → stem 兜底失配 → 残留警告而非谎报成功
            rb3 = run(KI, f_sp, "--kb", "b2", "--yes", "--build", "--json", env=henv)
            jload_str(rb3.stdout)
            (vb / "_conversion_report.json").unlink()
            rb4 = run(KI, "--rollback", "会议纪要 2026.md", "--kb", "b2", env=henv)
            check("高6: 无报告且 stem 失配 → 残留警告 + 不打成功行",
                  "警告" in (rb4.stderr or "") and "会议纪要-2026.md" in (rb4.stderr or "")
                  and "回滚不彻底" in (rb4.stderr or "") and "已回滚暂存" not in rb4.stdout,
                  (rb4.stderr or "")[:160])
            if doc_sp.exists():
                doc_sp.unlink()  # 清理残留，避免影响后续
            # 高6(c): 内容重复件回滚不误删 canonical
            (src17 / "正主.md").write_text("完全相同的内容X", encoding="utf-8")
            (src17 / "副本.md").write_text("完全相同的内容X", encoding="utf-8")
            rb5 = run(KI, src17 / "正主.md", src17 / "副本.md", "--kb", "b2", "--yes", "--build", "--json", env=henv)
            jload_str(rb5.stdout)
            dup17 = next((x for x in jload(vb / "_conversion_report.json")["details"]
                          if x["status"] == "skipped" and x.get("duplicate_of")), None)
            check("高6: build 产生 canonical+重复件", dup17 is not None)
            rb6 = run(KI, "--rollback", Path(dup17["source"]).name, "--kb", "b2", env=henv)
            check("高6: 重复件回滚保留 canonical",
                  Path(dup17["duplicate_of"]).exists() and "canonical" in rb6.stdout
                  and "警告" not in (rb6.stderr or ""), (rb6.stdout or "")[:120])
            # 中18: handoff 步骤 4 不再「可选」、指引重跑 finalize/emit
            f_h = src17 / "备注说明.md"; f_h.write_text("一段普通备注。", encoding="utf-8")
            rh = run(KI, f_h, "--kb", "b2", "--yes", env=henv)
            check("中18: handoff 摘要步指引重跑 finalize 且非可选",
                  "重跑 --finalize" in rh.stdout and "emit_access_bundle" in rh.stdout
                  and "4) 摘要（LLM·agent）" in rh.stdout and "可选" not in rh.stdout, (rh.stdout or "")[-200:])
            # 低13: finalize 双跑 use_when 不动点（不再自我回灌膨胀），且用户 extra 词保留
            run(KI, "--finalize", "--kb", "t", env=henv)
            uw1 = jload(av / "kb.json")["use_when"]
            run(KI, "--finalize", "--kb", "t", env=henv)
            uw2 = jload(av / "kb.json")["use_when"]
            check("低13: finalize 双跑 use_when 不动点 + extra 词保留",
                  uw1 == uw2 and "额外词" in uw1, uw1)

        with sect("[18] convert 边界: BOM / CSV 大字段 / 空文件 / 大小上限"):
            be = W / "be"; be.mkdir()
            (be / "bj.json").write_bytes(b"\xef\xbb\xbf" + json.dumps({"k": "v"}).encode())
            (be / "bt.txt").write_bytes(b"\xef\xbb\xbfBOM text")
            (be / "bm.md").write_bytes(b"\xef\xbb\xbf---\nname: x\n---\n\nbody")
            (be / "huge.csv").write_text("h1,h2\nA," + "Z" * 200000 + "\n", encoding="utf-8")
            (be / "big.txt").write_text("X" * 2_100_000, encoding="utf-8")
            (be / "biglist.json").write_text(json.dumps(list(range(1500))), encoding="utf-8")
            repe = W / "be.json"
            repe.write_text(json.dumps(report([
                (str(be / "bj.json"), "json"), (str(be / "bt.txt"), "text"), (str(be / "bm.md"), "markdown"),
                (str(be / "huge.csv"), "csv"), (str(be / "big.txt"), "text"), (str(be / "biglist.json"), "json"),
            ]), ensure_ascii=False), encoding="utf-8")
            ev = W / "ev"; run(SC / "convert_documents.py", "--scan-report", repe, "--output", ev)
            cre = jload(ev / "_conversion_report.json")
            dete = {Path(x["source"]).name: x for x in cre["details"]}
            check("中9: BOM JSON/text/md 全部转换成功",
                  all(dete[n]["status"] == "success" for n in ("bj.json", "bt.txt", "bm.md")),
                  {n: dete[n]["status"] for n in ("bj.json", "bt.txt", "bm.md")})
            bmd18 = (ev / "documents/bm.md").read_text(encoding="utf-8")
            check("中9: BOM 不再挡 front-matter 降级（单对 --- + 引用块）",
                  sum(1 for ln in bmd18.splitlines() if ln.strip() == "---") == 2 and "> name: x" in bmd18)
            hmd = (ev / "documents/huge.md").read_text(encoding="utf-8")
            check("中10: 超长单元格成功转换且截断标注",
                  dete["huge.csv"]["status"] == "success" and "单元格过长，已截断，原 200000 字符" in hmd
                  and len(hmd) < 10000, len(hmd))
            gmd = (ev / "documents/big.md").read_text(encoding="utf-8")
            check("中12: 超大 text 截断 + 标注",
                  dete["big.txt"]["status"] == "success" and "文件过大，仅转换前 2000000 字符" in gmd
                  and 2_000_000 <= len(gmd) <= 2_001_000, len(gmd))
            lmd = (ev / "documents/biglist.md").read_text(encoding="utf-8")
            check("中12: 超大 JSON 列表截断（保留 999、丢弃 1400）",
                  "共 1500 项，仅转换前 1000 项" in lmd and "- 999" in lmd and "- 1400" not in lmd)
            # 中11: 空白文件 → empty，不写盘、不登记
            ee = W / "ee"; ee.mkdir(); (ee / "empty.txt").write_text("   \n  \n", encoding="utf-8")
            repn = W / "ee.json"
            repn.write_text(json.dumps(report([(str(ee / "empty.txt"), "text")]), ensure_ascii=False), encoding="utf-8")
            ev2 = W / "ev2"; run(SC / "convert_documents.py", "--scan-report", repn, "--output", ev2)
            cre2 = jload(ev2 / "_conversion_report.json")
            check("中11: 空内容 → status=empty 且不写盘",
                  cre2["details"][0]["status"] == "empty" and "疑似扫描版" in cre2["details"][0]["reason"]
                  and not (ev2 / "documents/empty.md").exists() and cre2["empty"] == 1, cre2["details"][0])
            run(SC / "update_manifest.py", "--vault", ev2, "--from-conversion-report", ev2 / "_conversion_report.json")
            check("中11: empty 不登记 manifest（下一轮仍可重试）",
                  jload(ev2 / ".wiki-tree/manifest.json")["processed"] == {})

        with sect("[19] convert: mtime 沿管线 + source_path YAML 转义"):
            # 中13: 登记的 mtime = 扫描时刻而非登记时刻（抽取期间改文件不再被吞）
            mt = W / "mt"; mt.mkdir(); src_mt = mt / "mt.md"
            src_mt.write_text("内容M", encoding="utf-8")
            repm = W / "mt.json"
            repm.write_text(json.dumps(report([(str(src_mt), "markdown")],
                                              modified_at="2026-01-01T00:00:00+00:00"), ensure_ascii=False), encoding="utf-8")
            mv19 = W / "mv19"; run(SC / "convert_documents.py", "--scan-report", repm, "--output", mv19)
            crm = jload(mv19 / "_conversion_report.json")
            check("中13: 转换报告透传 modified_at",
                  crm["details"][0].get("modified_at") == "2026-01-01T00:00:00+00:00", crm["details"][0])
            os.utime(src_mt, (1_800_000_000, 1_800_000_000))  # 模拟抽取期间源被改
            run(SC / "update_manifest.py", "--vault", mv19, "--from-conversion-report", mv19 / "_conversion_report.json")
            recm = jload(mv19 / ".wiki-tree/manifest.json")["processed"][str(src_mt.resolve())]
            check("中13: 登记用扫描时刻 mtime（非登记时刻 stat）",
                  recm["mtime"] == "2026-01-01T00:00:00+00:00", recm["mtime"])
            run(SC / "scan_folder.py", mt, "--vault", mv19, "-o", W / "mt2.json")
            check("中13: 抽取期间的修改下一轮被识别为 modified",
                  jload(W / "mt2.json")["pending_count"] == 1)
            # 低3: source_path 含 " #" 写为 JSON 引号标量，YAML 不截断；同源重转无 -1
            yd = W / "y #dir"; yd.mkdir()
            wn = yd / "weird name.md"; wn.write_text("# w\nbody", encoding="utf-8")
            repy = W / "y.json"
            repy.write_text(json.dumps(report([(str(wn), "markdown")]), ensure_ascii=False), encoding="utf-8")
            yv = W / "yv"; run(SC / "convert_documents.py", "--scan-report", repy, "--output", yv)
            ymd = (yv / "documents/weird-name.md").read_text(encoding="utf-8")
            check("低3: source_path 为 JSON 引号标量整行",
                  f"source_path: {json.dumps(str(wn), ensure_ascii=False)}" in ymd, ymd[:200])
            run(SC / "convert_documents.py", "--scan-report", repy, "--output", yv)
            check("低3: 引号格式同源重转无 -1", not list((yv / "documents").glob("*-1.md")))
            # 低3: 旧裸值格式 front-matter 仍被识别为同源（兼容旧 vault）
            ld = W / "ld"; ld.mkdir(); lsrc = ld / "legacy.md"
            lsrc.write_text("# L\n正文", encoding="utf-8")
            ov = W / "ov"; (ov / "documents").mkdir(parents=True)
            (ov / "documents/legacy.md").write_text(
                f"---\nsource_type: local_file\nsource_path: {lsrc}\nsource_format: markdown\n---\n\n# L\n正文",
                encoding="utf-8")
            repl = W / "l.json"
            repl.write_text(json.dumps(report([(str(lsrc), "markdown")]), ensure_ascii=False), encoding="utf-8")
            run(SC / "convert_documents.py", "--scan-report", repl, "--output", ov)
            check("低3: 旧裸值行仍判同源（无 legacy-1.md）", not list((ov / "documents").glob("legacy-1.md")))

        with sect("[20] scan: 嵌套 vault 剪枝 + 内置剪枝可见化 + 孤儿"):
            # 高5: vault 嵌在扫描目标内 → 剪枝防自回灌
            nv = W / "nv"; nv.mkdir()
            (nv / "a.md").write_text("# A", encoding="utf-8")
            run(SC / "generate_wiki_structure.py", "--output", nv / "vault")
            (nv / "vault" / "documents" / "prev.md").write_text("# 上一轮产物", encoding="utf-8")
            run(SC / "scan_folder.py", nv, "--vault", nv / "vault", "-o", W / "nv.json")
            dnv = jload(W / "nv.json")
            check("高5: 嵌套 vault 被剪枝（vault_excluded + 产物不回灌）",
                  dnv["vault_excluded"] is True and "prev.md" not in {f["name"] for f in dnv["files"]}
                  and dnv["supported_files"] == 1, dnv.get("supported_files"))
            # 低5: 内置剪枝目录计入报告（相对 posix 路径、隐藏目录不进列表）
            bd = W / "bd"; (bd / "build" / "docs").mkdir(parents=True); (bd / "sub" / "dist").mkdir(parents=True)
            (bd / ".git").mkdir()
            (bd / "a.md").write_text("# A", encoding="utf-8")
            (bd / "build" / "docs" / "hid.md").write_text("x", encoding="utf-8")
            run(SC / "scan_folder.py", bd, "-o", W / "bd.json")
            dbd = jload(W / "bd.json")
            check("低5: builtin_excluded_dirs 可见（build/sub/dist，.git 不入列）",
                  dbd["builtin_excluded_dirs"] == ["build", "sub/dist"]
                  and "hid.md" not in {f["name"] for f in dbd["files"]}, dbd["builtin_excluded_dirs"])
            # 低17: 删除的源 → orphaned 可见；仅被 exclude 的不算孤儿
            ms2 = W / "ms2"; ms2.mkdir(); (ms2 / "x.md").write_text("内容X", encoding="utf-8")
            mvt2 = W / "mvt2"; run(SC / "generate_wiki_structure.py", "--output", mvt2)
            run(SC / "scan_folder.py", ms2, "-o", mvt2 / ".wiki-tree/scan.json")
            run(SC / "update_manifest.py", "--vault", mvt2, "--mark", str(ms2 / "x.md"), "--doc-md", "documents/x.md")
            run(SC / "scan_folder.py", ms2, "--vault", mvt2, "--exclude", "x.md", "-o", W / "or1.json")
            check("低17: 仅被 exclude（文件仍在盘上）不算孤儿", jload(W / "or1.json")["orphaned"] == [])
            (ms2 / "x.md").unlink()
            run(SC / "scan_folder.py", ms2, "--vault", mvt2, "-o", W / "or2.json")
            dor = jload(W / "or2.json")
            check("低17: 源文件删除 → orphaned 可见",
                  dor["orphaned_count"] == 1 and dor["orphaned"][0].endswith("x.md") and dor["pending_count"] == 0,
                  dor.get("orphaned"))
            (ms2 / "x.md").mkdir()  # 同名目录顶替源文件 → 源已非文件，仍应判孤儿
            run(SC / "scan_folder.py", ms2, "--vault", mvt2, "-o", W / "or3.json")
            check("低17: 同名目录顶替源文件仍判孤儿（isfile 判定）",
                  jload(W / "or3.json")["orphaned_count"] == 1)

        with sect("[21] reduce 健壮性: 坏类型 / 长实体名 / 文件名碰撞 / doclink / dedup 链"):
            # 中15: LLM 类型漂移不崩溃，计 bad_values
            b1v = W / "b1v"; run(SC / "generate_wiki_structure.py", "--output", b1v)
            (b1v / "documents/b1.md").write_text("# b1\nX Y", encoding="utf-8")
            (b1v / ".wiki-tree/extracted/b1.json").write_text(json.dumps({
                "doc_id": "b1", "doc_md": "documents/b1.md", "short_summary": None,
                "importance": "0.9", "topics": "主题串",
                "entities": [{"kind": "tool", "text": "X"}, {"kind": "tool", "text": "Y"}],
                "relations": [{"subject": "X", "predicate": "USES", "object": "Y", "confidence": "高"}]},
                ensure_ascii=False), encoding="utf-8")
            rb1v = run(SC / "assemble_vault.py", "--vault", b1v)
            ob1 = jload_str(rb1v.stdout)
            check("中15: 坏类型不崩溃（bad_values=2: confidence『高』+ topics 非列表）",
                  rb1v.returncode == 0 and ob1.get("bad_values") == 2 and (b1v / "_processing-report.md").exists(),
                  ob1.get("bad_values"))
            # 中16: 显式 --dedup-map 路径不存在 → 报错而非静默忽略
            rdm = run(SC / "assemble_vault.py", "--vault", b1v, "--dedup-map", W / "nope.json", ok=False)
            check("中16: 显式 dedup-map 缺失 → rc=1 + 报错", rdm.returncode == 1 and "不存在" in rdm.stderr,
                  (rdm.stderr or "")[:80])
            # 中17/低9/低10: 长实体名截断、casefold 碰撞、doclink 不信 LLM doc_id
            b2v = W / "b2v"; run(SC / "generate_wiki_structure.py", "--output", b2v)
            (b2v / "documents/doc2.md").write_text("# doc2\nAPI api", encoding="utf-8")
            longname = "超长实体名" * 30  # 150 字符
            (b2v / ".wiki-tree/extracted/j1.json").write_text(json.dumps({
                "doc_id": "WRONG-ID", "doc_md": "documents/doc2.md", "short_summary": "s", "importance": 0.5,
                "topics": ["t"],
                "entities": [{"kind": "tool", "text": "API"}, {"kind": "tool", "text": "api"},
                             {"kind": "tool", "text": longname}],
                "relations": [{"subject": "API", "predicate": "USES", "object": "api", "confidence": 0.9},
                              {"subject": "API", "predicate": "USES", "object": longname, "confidence": 0.8}]},
                ensure_ascii=False), encoding="utf-8")
            rb2v = run(SC / "assemble_vault.py", "--vault", b2v)
            ob2 = jload_str(rb2v.stdout)
            check("低10: doc_id 与 doc_md 不一致 → stderr 警告 + 采用推导值",
                  "WRONG-ID" in (rb2v.stderr or "")
                  and "[[doc2]]" in (b2v / "_index.md").read_text(encoding="utf-8")
                  and "[[WRONG-ID]]" not in (b2v / "_index.md").read_text(encoding="utf-8"), (rb2v.stderr or "")[:120])
            check("低9: casefold 碰撞分配后缀并报告",
                  len(ob2.get("collisions", [])) == 1 and ob2["collisions"][0]["file"] == "tool-api-2.md",
                  ob2.get("collisions"))
            long_cards = [p for p in (b2v / "entities").glob("tool-超长实体名*.md")]
            check("中17: 超长实体名截断建卡成功（errors 为空）",
                  ob2.get("errors") == [] and len(long_cards) == 1 and len(long_cards[0].stem) <= 100,
                  [p.name for p in long_cards])
            notes2 = {p.stem for p in b2v.rglob("*.md") if ".wiki-tree" not in p.parts}
            dangling2 = []
            for p in b2v.rglob("*.md"):
                if ".wiki-tree" in p.parts:
                    continue
                for m in re.finditer(r"\[\[([^\]]+)\]\]", p.read_text(encoding="utf-8")):
                    tgt = m.group(1).split("|")[0].split("#")[0].strip().split("/")[-1]
                    if tgt not in notes2:
                        dangling2.append((p.name, tgt))
            check("中17/低9: 截断+碰撞后仍零悬空 wikilink", not dangling2, str(dangling2[:5]))
            # 低7: dedup-map 链式解析 + 空值过滤 + 环检测（进程内调用）
            sys.path.insert(0, str(SC))
            import compute_centrality as _cc
            dmf = W / "chain.json"
            dmf.write_text(json.dumps({"A": "B", "B": "C", "X": "", "Y": None, "P": "Q", "Q": "P"},
                                      ensure_ascii=False), encoding="utf-8")
            err_buf = io.StringIO()
            with contextlib.redirect_stderr(err_buf):
                dmres = _cc.load_dedup_map(str(dmf))
            check("低7: 链式折叠 A→C、空值丢弃、环保持原值",
                  dmres == {"A": "C", "B": "C", "P": "Q", "Q": "P"}, dmres)
            check("低7: 空值/环有 stderr 警告", "空" in err_buf.getvalue() and "环" in err_buf.getvalue(),
                  err_buf.getvalue()[:120])

        with sect("[22] 文档契约守卫"):
            skill_md = (SKILL_ROOT / "SKILL.md").read_text(encoding="utf-8")
            check("中4: SKILL.md ≤16KB（原 28KB）", os.path.getsize(SKILL_ROOT / "SKILL.md") <= 16384,
                  os.path.getsize(SKILL_ROOT / "SKILL.md"))
            check("中5: when_to_use 字段已删除", "when_to_use:" not in skill_md)
            mdesc = re.search(r'^description: "(.*)"$', skill_md, re.M)
            check("中5: description ≤650 字符且含中英触发词",
                  bool(mdesc) and len(mdesc.group(1)) <= 650
                  and "第二大脑" in mdesc.group(1) and "Obsidian" in mdesc.group(1),
                  mdesc and len(mdesc.group(1)))
            check("高1: Phase 3 点名 convert_documents 且禁止手写",
                  "convert_documents.py --scan-report" in skill_md and "不要手写转换代码" in skill_md)
            check("中1: 点名 generate_wiki_structure", "generate_wiki_structure.py --output" in skill_md)
            check("高2: 增量流程点名重跑 emit", skill_md.count("emit_access_bundle.py") >= 3,
                  skill_md.count("emit_access_bundle.py"))
            check("中3: verify 用临时文件路径", ".wiki-tree/tmp/" in skill_md)
            ep = (SKILL_ROOT / "references/extraction-prompts.md").read_text(encoding="utf-8")
            check("中8: §4 含 one_liner 字段与落盘骨架",
                  '"one_liner"' in ep and "doc_count" in ep)
            check("中6: 关系端点逐字硬约束", "逐字使用实体列表中的 text" in ep)
            check("低2: JSON 硬约束全节覆盖", ep.count("严格返回 JSON") >= 5, ep.count("严格返回 JSON"))
            check("中7: topics 粒度约束", "1-3 个粗粒度领域词" in ep)
            sb = (SKILL_ROOT / "references/subagent-batch-extraction.md").read_text(encoding="utf-8")
            check("中2: 子 agent 契约带 skill_root",
                  "skill_root" in sb and "{skill_root}/references/extraction-prompts.md" in sb
                  and "{skill_root}/scripts/verify_entities.py" in sb)
            check("中3/中13: 契约用实体文件路径 + done 可带 mtime",
                  ".entities.json" in sb and '"mtime"' in sb)
            check("中6: 契约含关系 set 过滤兜底", "set 过滤" in sb)
            ab = SKILL_ROOT / "references/access-bundle.md"
            check("中4: access-bundle.md 存在且含下沉内容",
                  ab.exists() and all(k in ab.read_text(encoding="utf-8")
                                      for k in ("kb_hub_server", "kb_ingest", "emit_doc_summaries")))

        expect = os.environ.get("EXPECT_SKIPS")
        if expect is not None and len(skips) != int(expect):
            fails.append(f"skip 数 {len(skips)} != EXPECT_SKIPS {expect}（环境前置条件失效: {skips}）")
        print("\n" + "=" * 40)
        tail = f" ({len(skips)} skipped: {'; '.join(skips)})" if skips else ""
        print(("ALL PASS [OK]" + tail) if not fails else f"FAILED ({len(fails)}): {fails}" + tail)
        return 1 if fails else 0
    finally:
        shutil.rmtree(W, ignore_errors=True)


def test_skill_suite():
    """pytest 桥接：python -m pytest tests/test_skill.py 收集到本用例。"""
    assert main() == 0


if __name__ == "__main__":
    sys.exit(main())
